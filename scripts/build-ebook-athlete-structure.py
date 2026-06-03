#!/usr/bin/env python3
"""
Structure + premières sections rédigées de l'ebook Athlète Explosif
Sortie : ~/Library/CloudStorage/OneDrive-Personnel/Documents/Coaching/Ebook-Athlete-Explosif-Structure.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path.home()
OUTPUT_DIR = ROOT / "Library/CloudStorage/OneDrive-Personnel/Documents/Coaching"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUTPUT_DIR / "Ebook-Athlete-Explosif-Structure.docx"

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─── Titre principal ────────────────────────────────────────────────────────
title = doc.add_heading("Ebook Athlète Explosif", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Rayan Bonte · RB Perform · 2026")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
sub.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()

# ─── Helpers ────────────────────────────────────────────────────────────────
def add_h1(num, txt):
    h = doc.add_heading(f"{num}. {txt}", 1)
    return h

def add_h2(txt):
    h = doc.add_heading(txt, 2)
    return h

def add_para(txt):
    p = doc.add_paragraph(txt)
    return p

def add_note(txt):
    p = doc.add_paragraph()
    run = p.add_run(f"📝 {txt}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    return p

def add_bullet(txt):
    p = doc.add_paragraph(txt, style='List Bullet')
    return p

def add_space():
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# CHAPITRE 1 — PRÉAMBULE
# ═══════════════════════════════════════════════════════════════════════════
add_h1(1, "Préambule + Disclaimer légal")
add_para("Les informations contenues dans cet ebook sont fournies à titre éducatif et informatif uniquement.")
add_para("L'auteur n'est pas titulaire d'un diplôme d'État de coach sportif. Il ne propose ni encadrement sportif personnalisé, ni suivi individualisé, ni prescription médicale ou nutritionnelle.")
add_para("Toutes les méthodes présentées sont basées sur l'expérience personnelle de l'auteur en tant qu'athlète et sur les connaissances accumulées au contact de préparateurs physiques du sport de haut niveau.")
add_para("Avant de commencer ce programme, consulte un médecin si tu as des antécédents médicaux, des blessures ou des doutes sur ta capacité à suivre un entraînement intensif.")
add_para("Tous droits réservés. Toute reproduction est interdite sans l'autorisation de l'auteur.")
add_space()

# ═══════════════════════════════════════════════════════════════════════════
# CHAPITRE 2 — INTRO
# ═══════════════════════════════════════════════════════════════════════════
add_h1(2, "Intro")

add_h2("Bienvenue + parcours")
add_para("Bienvenue chez RB Perform.")
add_para("Je m'appelle Rayan Bonte. Rugbyman XIII semi-pro au SOA, athlète et passionné par la performance physique.")
add_para("Mon objectif avec cet ebook : t'accompagner vers la meilleure version de toi-même, avec des méthodes d'entraînement qui marchent vraiment.")
add_para("Le sport, je le pratique depuis toujours. J'ai eu la chance de côtoyer de nombreux athlètes de haut niveau et de tester sur le terrain les méthodes d'entraînement les plus avancées de la prépa physique moderne.")
add_para("Je tiens à le préciser dès maintenant : je ne suis pas coach sportif diplômé. Tous les conseils donnés ici ne sont en aucun cas des conseils médicaux. Ils sont uniquement basés sur mon expérience et mes connaissances d'athlète.")

add_para("📌 [Insère ici : Vidéo de présentation]").runs[0].italic = True

add_h2("Pourquoi tu es là")
add_para("Si tu lis ce programme, c'est que tu as fait le premier pas vers ton objectif. Le premier pas vers le changement.")
add_para("Tu es probablement un athlète qui veut être meilleur que son adversaire sur le terrain — et c'est ça qui doit te motiver chaque jour.")
add_para("Ou tu as juste de l'ambition. Tu veux te dépasser, devenir un meilleur athlète, jour après jour.")
add_para("Ou tu en as marre de ton physique, de tes habitudes actuelles, et tu te trouves pas assez performant.")
add_para("Quel que soit ton point de départ : tu es au bon endroit. Je vais t'expliquer ce qui fera la différence.")

add_h2("Premier exercice — fixe tes objectifs")
add_para("Premier conseil — et il vaut de l'or : fixe-toi des objectifs clairs. Sur 1 mois. Sur 3 mois. Sur 6 mois. Sur 1 an.")
add_para("L'humain est sans cesse à la recherche des hormones de bien-être qui se libèrent quand on réussit quelque chose (diplôme, compétition, record personnel). C'est ce qui te permettra de ne jamais lâcher et d'apprécier le processus jusqu'au bout.")
add_para("Chaque petite réussite compte. Perdre 500 g quand tu veux en perdre 10 te rapproche de ton objectif. Gagner 2 cm de vertical quand tu veux en gagner 10, pareil. Savoure chaque étape et tire-en encore plus de motivation.")
add_para("Maintenant, prends une feuille, un stylo, ou tes notes de téléphone. Écris au minimum 3 objectifs pour les échéances ci-dessus. Fais-le avant de continuer.")

add_h2("Ce que tu vas apprendre dans cet ebook")
add_para("Dans cet ebook, je vais te partager tout ce que j'ai accumulé pendant des années sur l'entraînement athlétique. Je vais t'éviter de perdre le temps que j'ai perdu — à chercher, à tester, à me tromper.")
add_para("Grâce à mon parcours, j'ai pu fréquenter ce qui se fait de mieux en préparation physique dans le sport de haut niveau. J'ai appris auprès de préparateurs physiques ultra-qualifiés. Tout ce qui marche est dans ces pages.")
add_para("À la fin de cet ebook, tu auras :")
add_bullet("Compris pourquoi un athlète ne s'entraîne pas comme un bodybuilder")
add_bullet("Maîtrisé les concepts clés : force, puissance, vitesse, pliométrie")
add_bullet("Un programme structuré de 8 semaines, complet (course + muscu + vidéos)")
add_bullet("Une routine de nutrition et compléments adaptée à un athlète")
add_bullet("Les clés de la mobilité et de la récupération pour durer sur une saison")
add_bullet("Le mindset qui sépare ceux qui finissent de ceux qui abandonnent")
add_para("Pas juste un programme. La compréhension complète de pourquoi tu fais chaque chose. Parce qu'un athlète qui sait ce qu'il fait, c'est un athlète qui progresse 2 fois plus vite.")
add_space()

# ═══════════════════════════════════════════════════════════════════════════
# CHAPITRE 3 — BODYBUILDER VS ATHLÈTE
# ═══════════════════════════════════════════════════════════════════════════
add_h1(3, "Bodybuilder vs Athlète")

add_para("Avant de parler de programme, faut clarifier un point essentiel : un athlète ne s'entraîne PAS comme un bodybuilder.")
add_para("C'est le piège n°1 des athlètes amateurs : ils copient les programmes qu'ils voient en salle, suivent les conseils des bodybuilders, et finissent forts en isolement... mais lents sur le terrain.")

add_h2("La différence : les objectifs")
add_para("Un pratiquant lambda de salle de sport s'entraîne pour son physique. Le reflet dans le miroir. Sa séance est construite pour isoler chaque muscle, le faire grossir, et trouver une harmonie visuelle. C'est parfait pour son objectif.")
add_para("Un athlète a deux objectifs très différents :")
add_bullet("Être performant sur le terrain (rapide, fort, explosif, endurant)")
add_bullet("Préparer son corps à encaisser la demande physique pour limiter les blessures et durer toute la saison")
add_para("Ces deux objectifs imposent des choix d'entraînement complètement différents.")

add_h2("Le problème de l'isolation")
add_para("Le bodybuilder se fout d'être fonctionnel. Il veut être beau. Il isole chaque muscle, les travaille séparément. Résultat : à long terme, problème de coordination entre les groupes musculaires.")
add_para("T'as déjà vu une vidéo d'un bodybuilder qui court ? Si oui, tu sais : ils sont raides, ça paraît pas naturel. Si non, va voir, c'est édifiant. Pourquoi ? Parce que leurs muscles n'ont jamais appris à travailler ensemble.")
add_para("L'entraînement athlétique cherche exactement l'inverse : optimiser la coordination musculaire et neuronale, pour être capable d'utiliser l'ensemble de son corps de la façon la plus efficace possible.")
add_para("C'est ça la première grande différence. Et c'est ça qui change tout sur le terrain.")

add_h2("Les pièges classiques à éviter")
add_para("Pendant des années, j'ai vu (et fait) les mêmes erreurs. Je te préviens d'avance.")

add_para("Piège n°1 : croire que plus tu es musclé, plus tu seras fort sur le terrain.")
add_para("C'est faux. Les films de superhéros nous ont vendu l'idée que les mecs les plus massifs sont toujours les meilleurs. Dans la réalité : la force et le poids comptent, oui, mais une force mal utilisée ne vaut rien sur le terrain. La technique, la coordination, la vitesse passent souvent avant.")

add_para("Piège n°2 : en faire trop. Trop de sprints, trop de vitesse, trop de muscu, en pensant que plus tu en fais, meilleur tu deviens.")
add_para("La progression vient du dosage et de la récupération, pas de l'accumulation.")

add_para("Piège n°3 : passer trop de temps en salle, au détriment de la technique et de la pratique du sport.")
add_para("La muscu, c'est un OUTIL au service du sport. Pas un sport en soi (sauf si tu fais du powerlifting). Si tu sors de salle épuisé et que tu loupes ton entraînement technique de rugby, t'as fait l'inverse de ce qui marche.")
add_para("Garde ces 3 pièges en tête. On va construire le programme pour les éviter.")

add_note("À DÉPLACER vers chapitre 4 (Théorie) : tout le contenu sur 'comment prendre du muscle', tension mécanique, stress métabolique, surcharge progressive. Et à DÉPLACER vers annexe technique : les méthodes avancées (rest-pause, séries dégressives, tempo, supersets, bisets).")
add_space()

# ═══════════════════════════════════════════════════════════════════════════
# CHAPITRES 4-15 — STRUCTURE À COMPLÉTER
# ═══════════════════════════════════════════════════════════════════════════

add_h1(4, "Théorie athlète")
add_para("Avant de te jeter dans le programme, faut que tu comprennes ce qui se passe sous le capot. Pas pour faire le malin avec des termes scientifiques — pour que tu saches pourquoi tu fais chaque chose. Un athlète qui comprend, c'est un athlète qui progresse 2 fois plus vite.")
add_para("Promesse : zéro jargon inutile. Juste les concepts qui changent vraiment ta façon de t'entraîner.")

add_h2("Force, puissance, vitesse — clarifions une fois pour toutes")
add_para("Ces trois mots sont mélangés tout le temps. C'est dommage parce qu'ils décrivent trois qualités différentes, et tu travailles chacune avec des méthodes différentes.")
add_para("La méthode la plus simple pour s'y retrouver : penser en types d'effort, pas en termes scientifiques.")
add_bullet("EFFORT MAXIMAL — tu soulèves très lourd (85-100% de ton 1RM) sur peu de répétitions (1-5). C'est ça qui construit ta force pure.")
add_bullet("EFFORT DYNAMIQUE — tu mobilises une charge légère à modérée (≤ 70% du 1RM) AVEC l'intention d'aller le plus vite possible. C'est ça qui construit ta puissance et ta vitesse.")
add_bullet("EFFORT RÉPÉTÉ — tu soulèves une charge moyenne (70-85%) sur beaucoup de répétitions jusqu'à très près de l'échec. C'est ça qui construit ta masse musculaire et ton endurance de force.")
add_para("La règle qui résume tout :")
add_para("« Quand t'es faible, deviens fort. Quand t'es fort, deviens rapide. Quand t'es rapide, apprends à répéter. »")
add_para("Un athlète sérieux doit faire les trois dans l'année. Pas tout en même temps. Dans cet ordre.")

add_h2("Le continuum force-vitesse — pourquoi tu ne peux pas tricher")
add_para("Voici une vérité physique que personne ne peut contourner :")
add_bullet("Plus la charge est lourde, plus tu la déplaces lentement.")
add_bullet("Plus la charge est légère, plus tu peux la déplacer vite.")
add_para("Entre les deux extrêmes, il existe un continuum : force max, puissance, vitesse, endurance. Pour un sport collectif comme le rugby, tu dois être bon partout sur ce continuum. Pas seulement fort. Pas seulement rapide. Les deux.")
add_para("C'est pour ça que ton programme doit varier les types d'effort dans la semaine. Une séance pour la force pure. Une séance pour la vitesse. Une séance pour la répétition. Pas trois fois la même chose.")

add_h2("La puissance — la qualité qui te sort du lot")
add_para("La puissance, c'est la capacité de produire de la force RAPIDEMENT. C'est THE qualité qui sépare un athlète de haut niveau d'un athlète moyen en sport collectif.")
add_para("Concrètement : deux joueurs qui squatent 120 kg, celui qui peut le faire plus vite est plus puissant. Et c'est lui qui sautera plus haut, sprintera plus vite, plaquera plus fort.")
add_para("Comment on travaille la puissance ?")
add_bullet("Mouvements balistiques : sauts (squat jump, box jump), épaulés, jetés, lancers de medecine ball. Pourquoi ? Parce que dans ces mouvements, tu ne décélères pas en fin de geste — tu projettes vraiment. C'est dans ces mouvements que tu peux exprimer ta puissance maximale.")
add_bullet("Intention de vitesse maximale : même si la charge n'est pas légère, tu pousses comme si elle l'était.")
add_para("Erreur classique : faire des squats à 60% du 1RM en mode \"je vais lentement pour faire propre\". Si l'objectif est la puissance, l'intention de vitesse est non-négociable.")

add_h2("Pliométrie — le levier sous-exploité")
add_para("La pliométrie, c'est l'entraînement basé sur des sauts et des bondissements. Pourquoi c'est puissant ?")
add_para("Quand tu atterris d'un saut, tes muscles s'étirent rapidement (phase excentrique). Pendant cet étirement, ils stockent de l'énergie élastique — comme un ressort comprimé. Si tu réagis immédiatement en re-sautant (phase concentrique), tu récupères cette énergie et tu sautes plus haut, plus vite, avec moins d'effort musculaire.")
add_para("C'est ce qu'on appelle le cycle étirement-raccourcissement. Et c'est la clé de l'explosivité.")
add_para("Concrètement :")
add_bullet("Box jumps — saut sur boîte. Travail vertical pur.")
add_bullet("Depth jumps — tu tombes d'un step, tu rebondis IMMÉDIATEMENT. Le truc qui fait vraiment progresser.")
add_bullet("Broad jumps — saut en longueur. Travail horizontal.")
add_bullet("Bondissements — sauts répétés sur place ou avec déplacement.")
add_para("Règles d'or de la pliométrie :")
add_bullet("Qualité > quantité. 3-5 reps par série, avec récup complète. Si tu sens la hauteur baisser, tu stop.")
add_bullet("Le temps au sol doit être COURT. Plus tu réagis vite, plus tu récupères d'énergie élastique.")
add_bullet("Pas tous les jours. 1-2 séances par semaine MAX. C'est très exigeant pour le système nerveux.")

add_h2("Le rack — ton outil principal en salle")
add_para("Le rack (ou cage à squat), c'est ta structure de sécurité quand tu fais des mouvements lourds avec barre. Comprendre comment l'utiliser correctement, c'est la différence entre une séance qui progresse et une séance qui te blesse.")
add_h2("Les types de racks que tu croiseras :")
add_bullet("Power-rack (cage complète) : 4 montants, sécurités réglables sur toute la hauteur. Le plus sûr pour travailler seul.")
add_bullet("Half-rack : 2 montants seulement, support de barre devant. Moins sécurisant mais prend moins de place.")
add_bullet("Squat stand : juste deux supports pour la barre. Pas de sécurité. Utilise un spotter ou ne va pas au max.")
add_para("Réglages essentiels avant chaque séance :")
add_bullet("Hauteur de barre : la barre doit être à la hauteur de tes clavicules quand tu te places sous (back squat) ou à hauteur d'épaule (front squat). Si tu dois te mettre sur la pointe des pieds pour la décrocher, c'est trop haut.")
add_bullet("Sécurités (les taquets en bas) : à régler 2-3 cm sous ta position basse. En cas de fail, la barre se pose sur les sécurités, pas sur toi.")
add_bullet("Position au pin : pour les fast lifts en partant du bas, place les taquets à la hauteur correspondant à la position de départ que tu veux travailler.")
add_para("Erreur classique : régler les sécurités trop bas \"pour pas qu'elles gênent\". Le jour où tu fail, la barre te broie le dos. 2-3 cm sous ta position basse, point.")

add_h2("Filières énergétiques — comment ton corps fabrique son carburant")
add_para("Ton corps a 3 voies principales pour fabriquer l'énergie de la contraction musculaire (l'ATP). Chacune correspond à un type d'effort :")
add_bullet("FILIÈRE ALACTIQUE (0 à 10 secondes) — la filière du sprint court, du saut, du placage. Très puissante, très brève. Aucune fatigue chimique.")
add_bullet("FILIÈRE LACTIQUE (10 secondes à 1-2 minutes) — la filière des efforts intenses prolongés. Tu produis des lactates (= la sensation de brûlure). C'est ce qu'on entraîne en sprints répétés (RSA).")
add_bullet("FILIÈRE AÉROBIE (au-delà de 2 minutes) — la filière de l'endurance. Le Bronco rugby. Le footing. Le tempo run.")
add_para("Point clé : les 3 filières travaillent en parallèle. C'est pas l'une OU l'autre. Mais selon la durée et l'intensité de l'effort, une filière domine.")
add_para("Pour un rugbyman, tu utilises TOUTES les filières dans un match (sprints courts, courses moyennes, efforts soutenus). C'est pour ça que ta préparation physique doit varier les formats : ne fais pas que du sprint, ne fais pas que du footing.")

add_h2("Comment travailler la vitesse — les volumes et formats")
add_para("La vitesse pure (= la capacité à courir vite), ça se travaille avec des règles précises. Sinon tu fais du cardio déguisé.")
add_para("Distances et volumes recommandés (par séance) :")
add_bullet("Accélérations 0-30m → 100 à 300 mètres total")
add_bullet("Vitesse max linéaire 40-60m → 100 à 300 mètres total")
add_bullet("Capacité anaérobie alactique 80-150m → double ces volumes")
add_bullet("Changements de direction 5-20m → divise par 2 (plus de stress musculaire)")
add_para("Si tu es en PLEINE SAISON, réduis tout à 30-100 mètres total. La fatigue d'entraînement ne doit pas dégrader ta performance en match.")
add_para("Récupération entre les sprints : 1 minute par 10 mètres de sprint. Sinon tu travailles pas la vitesse, tu travailles la résistance à la fatigue.")
add_bullet("Sprint 10m → 1 min récup")
add_bullet("Sprint 20m → 2 min récup")
add_bullet("Sprint 30m → 3 min récup")
add_bullet("Sprint 40m → 4 min récup")

add_h2("Le progrès dans l'ordre : Amplitude → Technique → Vitesse → Charge")
add_para("Quand tu apprends ou re-travaille un mouvement, suis cet ordre. Toujours.")
add_bullet("1. Amplitude complète — peux-tu faire le mouvement sur toute son amplitude ? Si non, travaille la mobilité d'abord.")
add_bullet("2. Technique parfaite — sur toute l'amplitude, sans charge ou avec barre vide.")
add_bullet("3. Vitesse — capacité à exécuter rapidement la phase concentrique.")
add_bullet("4. Charge — augmente le poids SEULEMENT quand 1, 2, 3 sont maîtrisés.")
add_para("Erreur n°1 des athlètes amateurs : ajouter de la charge avant d'avoir maîtrisé la technique. Résultat : technique qui se dégrade, plateau de progression, blessure inévitable.")

add_h2("Comment tu prends du muscle (les 3 leviers)")
add_para("Pour comprendre comment construire de la masse, retiens trois principes qui fonctionnent en parallèle :")
add_bullet("PROXIMITÉ DE L'ÉCHEC — être à 1-2 répétitions de l'échec sur tes séries d'hypertrophie. Sinon le stimulus est insuffisant.")
add_bullet("TENSION MÉCANIQUE — travailler avec des charges lourdes recrute un maximum de fibres musculaires. C'est la base.")
add_bullet("STRESS MÉTABOLIQUE — charges plus légères, séries longues près de l'échec. Crée le \"pump\" et stimule la croissance par voies hormonales.")
add_para("Les principes essentiels à respecter SI TU VEUX PROGRESSER :")
add_bullet("Technique propre — éviter les blessures et cibler vraiment le muscle.")
add_bullet("Surcharge progressive — c'est LE facteur n°1. Ajouter du volume, de la charge, ou des reps semaine après semaine.")
add_bullet("Temps de repos cohérent — pas plus de 4 minutes entre les séries de force. Plus court pour l'hypertrophie (1-2 min).")
add_bullet("Récupération nutritionnelle — protéines + glucides en quantité suffisante.")
add_bullet("Sommeil minimum 7-8h — non négociable. C'est pendant que tu dors que tu reconstruis.")
add_space()

add_h1(5, "Tests baseline")
add_para("Avant de commencer le programme, tu vas mesurer ton point de départ. Sans baseline, pas de progression visible. Et sans progression visible, pas de motivation à finir les 8 semaines.")
add_para("Voici les 5 tests à faire le Jour 1, puis à refaire à la fin de la Semaine 4 et de la Semaine 8.")

add_h2("Les 5 tests")
add_bullet("Sprint 10m — départ debout, chrono manuel ou appli (ex: Sprint Timer)")
add_bullet("Sprint 30m — départ debout, même chrono que le 10m")
add_bullet("Vertical jump — saut avec contre-mouvement, mesure avec tape ou appli (ex: My Jump Lab)")
add_bullet("1RM squat — ou estimation via 3-5 reps max si tu préfères pas tester à 100%")
add_bullet("1RM bench press — idem")
add_para("Optionnel selon ton sport :")
add_bullet("Bronco (rugby) — 1200m navette 60m en continu, chrono final")

add_h2("Tableau de progression — à remplir")

table = doc.add_table(rows=7, cols=5)
table.style = 'Light Grid Accent 1'
headers = ["Test", "Semaine 0", "Semaine 4", "Semaine 8", "Progression"]
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

rows_data = [
    "Sprint 10m (sec)",
    "Sprint 30m (sec)",
    "Vertical jump (cm)",
    "1RM squat (kg)",
    "1RM bench press (kg)",
    "Bronco / autre (sec)",
]
for i, label in enumerate(rows_data, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = label
    for j in range(1, 5):
        row_cells[j].text = ""

add_space()
add_note("Permet à l'acheteur de mesurer ses progrès sur les 8 semaines. Effet bonus : génère des témoignages naturels (\"j'ai gagné 5 cm de vertical en 8 semaines\").")
add_space()

add_h1(6, "Intro programme")
for item in [
    "Vidéo de présentation du programme",
    "Vue d'ensemble des 8 semaines",
    "Planning hebdomadaire type"
]:
    add_bullet(item)
add_space()

add_h1(7, "PROGRAMME 8 semaines")
for item in [
    "Séance course (vitesse, conditioning) — vidéos explicatives de chaque exo",
    "Séance musculation — vidéos explicatives de chaque exo",
    "Progression structurée semaine par semaine",
    "Cycle 1 (semaines 1-4) : construction des bases",
    "Cycle 2 (semaines 5-8) : intensification et pic de performance"
]:
    add_bullet(item)
add_space()

add_h1(8, "Nutrition athlète")
add_para("Le principe de base dans la recherche de performance, c'est que l'entraînement seul ne suffit pas. On prend souvent la métaphore de l'iceberg pour montrer tout ce qui se cache derrière un physique athlétique ou une performance hors norme. La nutrition fait partie de cette face immergée.")
add_para("Tu peux utiliser les méthodes d'entraînement les plus avancées de cet ebook — si tu ne respectes pas les principes nutritionnels de base, tes efforts seront vains.")
add_para("Tu connais probablement quelqu'un qui s'entraîne depuis des années mais qui n'a pas de résultats, uniquement parce qu'il ne surveille pas son alimentation. Beaucoup disent que l'alimentation représente 70% de tes résultats. Le chiffre exact se discute, l'idée non : sans nutrition adaptée, tu rates l'essentiel.")
add_para("Je vais te donner ici toutes les clés pour avoir une nutrition adaptée à tes objectifs d'athlète.")

add_h2("Pourquoi la nutrition d'athlète n'est pas la nutrition d'un pratiquant de salle")
add_para("Un pratiquant de salle mange pour son physique. L'athlète mange pour deux choses :")
add_bullet("Produire l'énergie nécessaire pour des entraînements intenses ET des matches")
add_bullet("Récupérer plus vite, pour pouvoir enchaîner les séances et durer toute la saison")
add_para("C'est cette double exigence qui change l'approche. Tu ne peux pas te permettre la même rigidité qu'un bodybuilder en prep, parce qu'un athlète a besoin de glucides en abondance, d'énergie disponible, de fluidité. Mais tu ne peux pas non plus manger au feeling — sinon ta progression stagne.")

add_h2("Casser un mythe : transformer le gras en muscle, c'est impossible")
add_para("Je préfère le dire tout de suite : le gras et le muscle sont deux tissus complètement différents. On ne peut pas \"convertir\" l'un en l'autre.")
add_para("MAIS : durant les premières années d'entraînement, il est possible de perdre du gras ET gagner du muscle simultanément, à condition que ton déficit calorique reste léger (5-10%). Ça s'appelle la recomposition corporelle. C'est une zone réservée aux débutants/intermédiaires.")
add_para("Au-delà, il faut choisir une phase à la fois : soit prise de masse, soit perte de gras.")

add_h2("Étape 1 — Calculer ton métabolisme de base (TMB)")
add_para("Le TMB, c'est le nombre de calories que ton corps brûle au repos, juste pour fonctionner (respiration, digestion, cerveau, etc.). C'est ton point de départ.")
add_para("Deux formules à utiliser. La première si tu ne connais pas ton taux de masse grasse, la deuxième si tu le connais (plus précise).")

add_para("Formule Harris-Benedict (sans masse grasse)")
add_bullet("Hommes : 66,4730 + (13,7516 × poids en kg) + (5,0033 × taille en cm) − (6,7550 × âge en années)")
add_bullet("Femmes : 655,0955 + (9,5634 × poids en kg) + (1,8496 × taille en cm) − (4,6756 × âge en années)")

add_para("Formule Katch-McArdle (avec masse grasse — plus précise)")
add_para("D'abord, calcule ta masse maigre :")
add_para("Masse maigre = Poids total − (Poids total × % masse grasse)")
add_para("Puis applique :")
add_para("TMB = Masse maigre × 21,6 + 370")
add_para("Cette deuxième formule est nettement plus précise si tu connais ton taux de gras (mesure avec impedance, pince ou estimation visuelle).")

add_h2("Étape 2 — Prendre en compte ton niveau d'activité")
add_para("Le TMB seul ne suffit pas. Tu dépenses plus de calories que ça, surtout en tant qu'athlète. Multiplie ton TMB par un facteur selon ton niveau réel :")
add_bullet("Sédentaire (peu/pas d'exercice) : TMB × 1,2")
add_bullet("Légèrement actif (1-3 séances/semaine) : TMB × 1,35")
add_bullet("Assez actif (6-7 séances/semaine) : TMB × 1,5")
add_bullet("Très actif (exercice intense quotidien) : TMB × 1,7")
add_bullet("Extrêmement actif (2+ séances intenses/jour) : TMB × 1,9")
add_para("Le résultat = ton estimation de dépense calorique quotidienne (TDEE). C'est ce que tu dois consommer pour maintenir ton poids actuel.")
add_para("C'est une valeur théorique. Ajuste-la sur 2-3 semaines en fonction de ce que tu vois sur la balance et au miroir.")

add_h2("Étape 3 — Choisir ton objectif")
add_para("Une fois ton TDEE connu, tu décides où tu te places.")

add_para("Prise de masse")
add_bullet("Surplus calorique de 0 à 10%")
add_bullet("Croissance musculaire optimale en limitant la prise de gras")
add_bullet("Note honnête : lors d'une prise de masse, un peu de gras est presque obligatoire. Le but est de la limiter, pas de l'éviter.")

add_para("Perte de gras durable")
add_bullet("Déficit calorique de 15 à 20%")
add_bullet("Rythme soutenu mais durable, compatible avec ton entraînement")

add_para("Recomposition corporelle")
add_bullet("Déficit léger de 5 à 10%")
add_bullet("Pour construire du muscle ET éliminer la graisse simultanément")
add_bullet("Réservé aux débutants/intermédiaires (au-delà, plus dur)")

add_para("Perte de gras rapide (urgence)")
add_bullet("Déficit calorique de 20% ou plus")
add_bullet("Solution pour une courte période uniquement (4-6 semaines max)")
add_bullet("Risque : perte de masse musculaire, baisse d'énergie, fatigue")

add_para("BONUS ATHLÈTE — In-season vs Off-season")
add_bullet("En SAISON : tu maintiens ou tu prends légèrement (surplus 0-5%). Tu as besoin d'énergie pour les matches. Pas le moment de te restreindre.")
add_bullet("En HORS-SAISON (préparation) : tu peux te permettre des phases plus marquées (déficit 15-20% pour sécher, surplus 5-10% pour prendre de la masse).")

add_h2("Étape 4 — Répartir les macronutriments")

add_para("Les protéines — la fondation")
add_para("Quand on pense à une alimentation sportive, la première chose qui vient à l'esprit, c'est qu'il faut consommer des protéines. Voici à quoi elles servent vraiment.")
add_bullet("RÔLE : réparer et reconstruire les tissus musculaires après l'entraînement intensif.")
add_bullet("APPORT : 1 g de protéine = 4 calories.")
add_bullet("RECOMMANDATION : 2,2 à 2,5 g par kg de masse maigre pour une récupération optimale.")
add_bullet("SOURCES : viandes, poissons, œufs, légumineuses, produits laitiers, protéines végétales.")

add_para("Les glucides — le carburant de performance")
add_para("Pour moi, les glucides sont LA variable alimentaire la plus importante pour la performance sportive. La quantité que tu consommes détermine l'intensité de l'effort que tes muscles pourront fournir.")
add_para("Il existe 3 formes de glucides (monosaccharides, disaccharides, polysaccharides), mais pas besoin de rentrer dans le détail ici. Retiens juste l'essentiel :")
add_bullet("RÔLE : fournir l'énergie immédiate aux muscles pendant l'effort.")
add_bullet("APPORT : 1 g de glucide = 4 calories.")
add_bullet("RECOMMANDATION : représentent le reste de tes calories une fois protéines et lipides fixés.")
add_bullet("SOURCES : féculents (riz, pâtes, quinoa), tubercules (pommes de terre, patates douces), légumineuses (lentilles, haricots), fruits et légumes.")

add_para("Les lipides — santé et énergie longue durée")
add_para("Les graisses ont eu mauvaise image pendant des années. La réalité : c'est ta deuxième source d'énergie après les glucides, et c'est essentiel pour ton équilibre hormonal et neurologique.")
add_bullet("RÔLE : réserve d'énergie concentrée, structure cellulaire, santé du cerveau, mémoire, concentration, équilibre hormonal.")
add_bullet("APPORT : 1 g de lipide = 9 calories (plus du double des protéines et glucides).")
add_bullet("DOSAGE MINIMUM : 0,8 g par kg de poids de corps pour maintenir les fonctions hormonales.")
add_bullet("SOURCES : huiles végétales, noix, amandes, avocats, poissons gras (saumon, maquereau), graines.")
add_para("Deux catégories : graisses saturées et insaturées. Privilégie les insaturées sans diaboliser les saturées — les deux ont leur rôle.")

add_para("Les légumes — la base trop souvent négligée")
add_para("Faibles en calories, denses en vitamines, minéraux et fibres. Tu peux les consommer en grande quantité, varie les types pour couvrir tous les nutriments.")
add_bullet("RÔLE : satiété grâce aux fibres, apport en vitamines et minéraux essentiels à la récupération.")
add_bullet("SOURCES : légumes verts (brocoli, épinard), légumes secs (lentilles, pois chiches), légumes racines (carottes, patates douces).")

add_h2("Récap : la méthode pour répartir tes macros")
add_bullet("1. Détermine si tu veux un déficit ou un surplus calorique selon ton objectif.")
add_bullet("2. Fixe tes protéines : 2,2 à 2,5 g par kg de masse maigre.")
add_bullet("3. Fixe tes lipides : 0,8 à 1 g par kg de poids de corps.")
add_bullet("4. Le reste des calories = glucides (= ton carburant performance).")

add_h2("Étape 5 — Timing autour de l'entraînement")

add_para("Avant l'entraînement (1-2 h avant)")
add_para("Objectif : optimiser ton énergie pour la séance.")
add_bullet("Source de glucides à assimilation rapide : riz blanc, flocons d'avoine, pain blanc, fruits, compote.")
add_bullet("Petite portion de protéines maigres si tu peux digérer.")
add_bullet("Éviter les lipides en grande quantité (digestion plus lente, lourdeur).")

add_para("Pendant l'entraînement (si > 1h ou intensité élevée)")
add_bullet("Boisson énergétique (eau + sucre rapide + sels minéraux).")
add_bullet("Ou alternative simple : bonbons style Tagada, dattes, bananes coupées.")
add_para("Pas obligatoire pour une séance courte de 45-60 min, mais utile sur les blocs longs ou doubles séances.")

add_para("Après l'entraînement (dans l'heure qui suit)")
add_para("Objectif : reconstituer les réserves de glycogène musculaire et démarrer la reconstruction musculaire.")
add_bullet("Protéines pour la récupération (whey si tu en prends, sinon poulet/œufs/skyr).")
add_bullet("Glucides rapides (banane, miel, riz blanc).")

add_h2("Étape 6 — Timing autour du match (spécifique athlète)")

add_para("J-1 (la veille au soir)")
add_bullet("Repas riche en glucides complexes (pâtes, riz, patates douces).")
add_bullet("Hydratation maximale.")
add_bullet("Pas d'aliments nouveaux ou risqués (digestion).")

add_para("Jour J — 3 h avant le match")
add_bullet("Repas complet : glucides lents + protéines maigres + peu de lipides.")
add_bullet("Exemple : riz blanc + poulet + légumes verts.")

add_para("Jour J — 30-60 min avant le match")
add_bullet("Glucides rapides légers : banane, compote, gâteau de riz, dattes.")
add_bullet("Hydratation finale (300-500 ml d'eau).")

add_para("Après le match")
add_bullet("Dans l'heure : protéines + glucides rapides pour recharger le glycogène.")
add_bullet("Repas complet 2-3 h après pour la récupération longue.")

add_h2("Étape 7 — Exemple de journée type (2800 kcal)")
add_para("Pour un athlète de 80-85 kg en maintien ou léger surplus. Cible : 150 g protéines / 80 g lipides / 370 g glucides.")

add_para("Petit déjeuner")
add_bullet("50 g flocons d'avoine")
add_bullet("250 ml lait demi-écrémé")
add_bullet("1 banane")
add_bullet("20 g beurre de cacahuète")

add_para("Déjeuner")
add_bullet("150 g riz basmati cru")
add_bullet("150 g blanc de poulet")
add_bullet("10 g huile d'olive")
add_bullet("Légumes verts à volonté")

add_para("Collation")
add_bullet("100 g pain complet")
add_bullet("20 g miel")
add_bullet("30 g amandes")
add_bullet("150 g skyr")

add_para("Dîner")
add_bullet("150 g pâtes complètes crues")
add_bullet("150 g steak 5%")
add_bullet("10 g huile d'olive")
add_bullet("Légumes à volonté")

add_para("Variantes à construire selon ton profil :")
add_bullet("Athlète 75 kg, perte de gras → cible ~2200-2400 kcal")
add_bullet("Athlète 85 kg, maintien → ~2800 kcal (exemple ci-dessus)")
add_bullet("Athlète 95+ kg, prise de masse → ~3200-3500 kcal")

add_h2("Étape 8 — Peut-on se faire plaisir ?")
add_para("Pour qu'une transformation soit réussie et durable, le régime ne doit pas être une contrainte trop importante.")
add_para("Des repas plaisir peuvent tout à fait faire partie de ta diète, à condition d'adapter les autres repas pour respecter ton total calorique.")
add_para("Concernant le choix des aliments, tu es assez libre, tant que tu respectes tes macronutriments et ton total calorique. Manger des aliments que tu aimes est essentiel pour tenir sur la durée.")

add_h2("La règle simple pour savoir si un aliment est sain")
add_para("« Est-ce que je peux le trouver dans la nature ? »")
add_para("Si la réponse est oui, l'aliment est relativement sain. Si la réponse est non, c'est probablement un produit transformé — pas interdit, mais à consommer en quantité modérée.")

add_note("À COMPLÉTER : liste d'aliments sains avec équivalents (pour varier) + 3-5 recettes (dont 1-2 végétariennes) + tableau d'aliments transformés OK vs à éviter.")
add_space()

add_h1(9, "Compléments alimentaires + électrolytes")

add_h2("Le principe avant tout")
add_para("Les compléments alimentaires ne sont pas des solutions miracles. Ni une forme de dopage. Ils sont là pour t'aider à combler des carences qu'une alimentation saine ne permet pas toujours de combler, et pour optimiser certaines fonctions précises.")
add_para("Mais ils ne remplaceront jamais une alimentation riche et variée. Si ta base nutritionnelle est cassée, prendre 10 produits ne sauvera rien.")
add_para("Je vais te parler ici uniquement des compléments qui valent vraiment le coup pour un athlète. Le reste, c'est du marketing.")

add_h2("Les 5 compléments qui marchent vraiment")

add_para("1. Whey protéine")
add_para("La whey, c'est une forme de protéine en poudre, exactement comme dans la viande, le poisson ou les œufs. Elle vient du liquide séparé du lait lors de la fabrication du fromage. Aucun danger contrairement à ce que beaucoup pensent.")
add_para("Ses avantages :")
add_bullet("Assimilation rapide après l'entraînement")
add_bullet("Pratique (transportable, rapide à préparer)")
add_bullet("Économique par gramme de protéine vs viande/poisson")
add_para("Mon avis honnête : la whey n'est pas indispensable. Un déficit en protéines peut être comblé en mangeant plus de viande, poisson ou œufs. Mais en avoir chez toi est utile quand tu manques de temps ou que tu pars en déplacement.")
add_para("Conseil pratique : si tu en consommes, le meilleur moment est immédiatement après l'entraînement. La quantité dépend de tes apports protéiques totaux sur la journée.")

add_para("2. Créatine")
add_para("La créatine, c'est une molécule présente naturellement en faible quantité dans la viande et le poisson. Dans le muscle, elle se transforme en phosphocréatine qui aide à produire rapidement de l'énergie pour les efforts courts et intenses : sprints, sauts, séries lourdes.")
add_para("Pourquoi tu devrais en prendre :")
add_bullet("Augmentation de la force et de la puissance maximale")
add_bullet("Capacité à maintenir des efforts intenses plus longtemps")
add_bullet("Gain de masse musculaire (en partie via hydratation cellulaire)")
add_bullet("Favorise la synthèse protéique")
add_para("C'est probablement LE complément avec le plus de preuves scientifiques. Pour un athlète explosif, c'est non négociable.")
add_para("Conseil pratique : 3 à 10 g par jour selon ta masse musculaire, réparti sur la journée. À consommer avec des glucides pour optimiser l'absorption. Pas besoin de phase de charge, prends ta dose quotidienne en continu.")

add_para("3. Oméga-3")
add_para("Les Oméga-3 (huile de poisson) sont des acides gras essentiels composés d'EPA et de DHA. Ton corps ne peut pas les fabriquer, tu dois les apporter par l'alimentation ou la supplémentation.")
add_para("Ils jouent un rôle clé dans :")
add_bullet("Le bon fonctionnement du cœur et du cerveau")
add_bullet("La régulation hormonale")
add_bullet("L'inflammation (propriétés anti-inflammatoires)")
add_bullet("La récupération musculaire")
add_bullet("La synthèse des protéines")
add_para("Le problème : la majorité des gens ne consomment pas assez de poissons gras (saumon, maquereau, sardines) pour couvrir leurs besoins. La supplémentation est souvent justifiée.")
add_para("Conseil pratique : 2 à 3 g par jour, possible d'aller jusqu'à 8 g selon les besoins.")

add_para("4. Magnésium")
add_para("76% des Français sont déficitaires en magnésium (source : ANSES). Pour un athlète qui sue, ce chiffre monte encore. Et la carence en magnésium se voit directement : paupières qui sautent, crampes, fatigue, sommeil de mauvaise qualité, irritabilité, tensions musculaires.")
add_para("Le magnésium intervient dans plus de 300 réactions biochimiques de ton corps. Il joue un rôle critique pour :")
add_bullet("La contraction musculaire (= TES performances)")
add_bullet("La transmission nerveuse")
add_bullet("La production d'ATP (= ton carburant)")
add_bullet("La gestion du stress et l'anxiété")
add_bullet("La qualité du sommeil")
add_para("Les formes à privilégier (bien absorbées) :")
add_bullet("Bisglycinate (effet calmant, idéal pour le soir)")
add_bullet("Citrate (recommandé en cas de constipation)")
add_bullet("Malate (forme énergisante, idéal pour le matin)")
add_bullet("Taurinate (agit au niveau cérébral)")
add_para("Les formes à éviter (mal absorbées, inconforts digestifs) :")
add_bullet("Oxyde, sulfate, chlorure, carbonate, lactate, magnésium marin")
add_para("Conseil pratique :")
add_bullet("Dose pour athlète : 500 mg/jour de magnésium élémentaire (vérifie le pourcentage VNR sur l'étiquette)")
add_bullet("Pour le sommeil/récupération : à prendre le soir")
add_bullet("Pour le stress : matin ET soir")
add_bullet("Cycle : 3 semaines puis 1 semaine de pause OU 3 mois puis 1 mois de pause")
add_bullet("⚠️ Ne pas prendre en cas d'insuffisance rénale")
add_bullet("⚠️ À distance du fer et du calcium (3-4 h)")

add_para("5. Vitamine D")
add_para("La vitamine D est essentielle au bon fonctionnement de tes systèmes nerveux, musculaire, immunitaire et osseux.")
add_para("Le problème en France : on est tous plus ou moins carencés, surtout d'octobre à mars (manque d'exposition solaire). Et un athlète carencé en vitamine D, c'est un athlète qui récupère mal, dort mal, et qui tombe malade plus souvent.")
add_para("Conseil pratique : 2000 à 4000 UI par jour selon la saison et l'exposition solaire. Idéalement combiné avec de la vitamine K2 (favorise la fixation osseuse).")

add_h2("Bonus athlète — Électrolytes & hydratation")

add_para("« L'eau ne suffit pas à t'hydrater. En fait, boire trop d'eau sans les minéraux appropriés peut te déshydrater encore plus. »")
add_para("C'est le principe central que la plupart des athlètes amateurs ignorent. L'hydratation ne se résume pas à boire de l'eau. Sans minéraux (électrolytes), l'eau ne pénètre pas correctement les cellules.")

add_para("Pourquoi tu es probablement carencé")
add_bullet("Le stress moderne appauvrit les minéraux du corps")
add_bullet("Les aliments d'aujourd'hui sont moins riches en nutriments qu'avant")
add_bullet("L'exercice physique intensif augmente les besoins")
add_bullet("L'eau du robinet, en bouteille ou filtrée est souvent dépourvue de minéraux")

add_para("Les 4 électrolytes principaux à connaître")
add_bullet("Sodium (Na) : équilibre des fluides, transmission nerveuse, contraction musculaire. Le discours \"anti-sel\" est exagéré pour un athlète qui sue. Tu en as besoin.")
add_bullet("Potassium (K) : contraction musculaire, équilibre acido-basique, rythme cardiaque. Perte importante pendant l'exercice intense.")
add_bullet("Magnésium (Mg) : 300+ réactions biochimiques, contraction musculaire, ATP, sommeil.")
add_bullet("Chlorure (Cl) : équilibre des liquides corporels, pression sanguine.")

add_para("Combien d'eau boire ?")
add_para("Repère : 1 litre d'eau de qualité pour 22 kg de poids de corps quand tu es très actif.")
add_bullet("Athlète 75 kg actif → ~3,4 L/jour")
add_bullet("Athlète 85 kg actif → ~3,9 L/jour")
add_bullet("Athlète 95 kg actif → ~4,3 L/jour")
add_para("Tu commences déjà ta journée déshydraté : pendant la nuit tu perds environ 1 kg de poids de corps (presque exclusivement eau + électrolytes). Bois 500 ml dès le réveil.")
add_para("Pendant un match intense (rugby, foot, basket), tu peux perdre 2 à 4 kg majoritairement d'eau. Conséquences : moins de lucidité, moins d'énergie, crampes, récupération dégradée.")

add_para("La règle pratique simple pour t'hydrater vraiment")
add_para("1 pincée de gros sel naturel (Celtique ou Himalaya, PAS le sel de table) dans 500 ml d'eau.")
add_para("Pourquoi ça marche : le sel naturel aide ton corps à absorber l'eau et à rester hydraté plus longtemps. Tes cellules dépendent de ces minéraux pour fonctionner.")
add_para("Jette le sel de table industriel. Investis dans du vrai sel.")

add_para("Apports recommandés par activité")
add_bullet("Entraînement < 2 h : ton alimentation couvre tes pertes. Pas besoin de boisson d'effort complexe.")
add_bullet("Entraînement > 2 h ou forte chaleur : boisson d'effort maison (eau + sel naturel + jus de fruit ou fruits frais).")
add_bullet("Apport potassium pendant l'effort long : 200 à 400 mg par heure d'effort.")
add_bullet("Pendant un match : boire toutes les 15-20 min, pas en grande quantité d'un coup.")

add_para("Sources alimentaires d'électrolytes")
add_bullet("Riches en potassium : banane, pomme de terre, patate douce, épinard, avocat, cacao cru, fruits secs, légumineuses")
add_bullet("Riches en magnésium : chocolat noir 70%+, amandes, graines de courge, quinoa, sardines, épinards, avocat, cajou")
add_bullet("Sources naturelles d'électrolytes : eau de coco, pastèque, fraise, orange, banane, tomate, brocoli, olive, poisson")

add_h2("Les compléments overrated (économise ton argent)")

add_para("Ce que tu peux laisser tomber sans regret :")
add_bullet("BCAA en poudre : si tu manges assez de protéines (2 g/kg masse maigre), les BCAA en plus apportent peu. Marketing intense, ROI faible.")
add_bullet("Glutamine : utile uniquement dans des cas très spécifiques (post-chirurgie, immunodépression). Pour un athlète sain, l'alimentation suffit.")
add_bullet("Pré-workout complexe (15 ingrédients) : ce qui marche dedans, c'est la caféine (200-400 mg) + la bêta-alanine (3-6 g). Tu peux les acheter séparément pour 10× moins cher.")
add_bullet("Brûleurs de graisse : la perte de gras se fait par le déficit calorique. Aucun complément ne va remplacer ça. Souvent dangereux pour le cœur.")
add_bullet("HMB, ZMA, tribulus, testo boosters : effets marginaux à nuls scientifiquement.")

add_h2("Récapitulatif — la stack basique de l'athlète")
add_para("Si tu veux une routine simple et efficace, voici ce que je te conseille comme base :")
add_bullet("Créatine monohydrate : 3-5 g/jour, en continu")
add_bullet("Oméga-3 : 2-3 g/jour")
add_bullet("Magnésium bisglycinate : 500 mg/jour le soir (cycles 3 sem ON / 1 sem OFF)")
add_bullet("Vitamine D3 + K2 : 2000-4000 UI/jour d'octobre à mars")
add_bullet("Sel naturel (Celtique/Himalaya) : 1 pincée dans tes 500 ml d'eau quotidiens")
add_bullet("Whey : optionnelle, selon ton besoin pratique")
add_para("Coût mensuel total : 30 à 50 € selon les marques. Bien moins que tu ne dépenses probablement en pré-workouts marketing.")

add_h2("Ce qui te fera plus progresser que TOUS ces compléments")
add_para("Avant d'acheter un seul produit, vérifie que tu coches ces 3 cases :")
add_bullet("Tu dors au moins 7 heures par nuit")
add_bullet("Tu manges 3-4 vrais repas par jour avec assez de protéines")
add_bullet("Tu t'hydrates correctement (eau + sel naturel)")
add_para("Si tu rates une de ces 3 cases, aucune supplémentation ne sauvera tes performances. C'est dans cet ordre que tu attaques.")
add_space()

add_h1(10, "Mobilité / Étirements")

add_h2("Pourquoi la mobilité n'est pas optionnelle pour un athlète")
add_para("La mobilité, c'est la capacité de ton corps à bouger dans toute son amplitude sans restriction. C'est différent de la flexibilité passive : la mobilité, c'est actif, c'est ce que tu peux faire toi-même sous contrôle.")
add_para("Pour un athlète, c'est non négociable. Trois raisons :")
add_bullet("PERFORMANCE — Tu ne peux pas sauter haut si tes chevilles sont raides. Tu ne peux pas squatter profond si tes hanches sont bloquées. Tu ne peux pas plaquer correctement si tes épaules sont fermées.")
add_bullet("PRÉVENTION BLESSURES — Un muscle qui ne peut pas s'étirer va finir par se déchirer. La plupart des blessures musculaires arrivent en bout d'amplitude.")
add_bullet("RÉCUPÉRATION — Tissus mobiles = meilleure circulation sanguine = meilleure récupération entre les séances.")
add_para("Erreur n°1 des athlètes amateurs : zéro travail de mobilité. Ils s'entraînent dur, ils ne s'étirent jamais. Résultat : ils plafonnent en performance et ils se blessent.")
add_para("Bonne nouvelle : 10 minutes par jour suffisent.")

add_h2("Les 3 zones critiques à travailler")

add_para("1. Les hanches")
add_para("La zone la plus importante pour un athlète. Tes hanches dictent ta capacité à sprinter, sauter, squatter, changer de direction. Si elles sont raides, tout le reste suit.")
add_para("Exercices clés :")
add_bullet("90/90 : assis au sol, jambe avant à 90°, jambe arrière à 90°. Tourne d'un côté à l'autre lentement. 5 reps par côté.")
add_bullet("Pigeon stretch : hanche en rotation externe, l'autre jambe étendue derrière. Tiens 30-60s par côté.")
add_bullet("Hip flexor stretch : fente avant, pousse les hanches en avant, lève le bras opposé. 30s par côté.")
add_bullet("Hip CARs (Controlled Articular Rotations) : à 4 pattes, lève une jambe et fais des cercles lents avec le genou. 5 reps par côté.")
add_para("Volume : 5-10 min par jour. Idéalement le matin ou après l'entraînement.")

add_para("2. Les chevilles")
add_para("Critique pour le squat profond (descente sous parallèle), le sprint et les changements de direction. Si tes chevilles ne fléchissent pas assez, tu compenses avec le dos = blessure garantie à terme.")
add_para("Exercices clés :")
add_bullet("Knee to wall test : à genou contre un mur, le genou avant doit toucher le mur sans que le talon décolle. Vise 10-12 cm entre les orteils et le mur. Travaille des 2 côtés.")
add_bullet("Calf raises + descente lente : monte sur la pointe des pieds, descends lentement en allant en-dessous du niveau du sol (sur une marche). 3 séries de 10-15.")
add_bullet("Ankle CARs : assis sur une chaise, dessine des cercles lents avec ton pied. 10 cercles par côté.")
add_para("Volume : 5 min par jour suffit, mais à faire en continu.")

add_para("3. Les épaules + colonne thoracique")
add_para("Pour les développés, les tractions, les overhead lifts, les plaquages, les passes au rugby. Une épaule qui ne s'ouvre pas = limitations partout et risque de blessure haute.")
add_para("Exercices clés :")
add_bullet("Shoulder dislocates avec bâton : prends un balai ou un bâton, mains larges, passe-le devant toi puis derrière toi en gardant les bras tendus. 10 reps. Rapproche les mains au fur et à mesure que tu progresses.")
add_bullet("Wall slides : dos au mur, bras en U contre le mur. Glisse les bras vers le haut puis vers le bas en gardant le contact avec le mur. 10 reps.")
add_bullet("T-spine rotation : à 4 pattes, main derrière la tête. Tourne le coude vers le plafond, puis vers le sol. 10 reps par côté.")
add_para("Volume : 5-10 min par jour.")

add_h2("Routine express 10 minutes — à faire quotidiennement")
add_para("Si tu n'as pas le temps de faire toute la routine en détail, voici une version condensée qui te prend 10 minutes et qui couvre l'essentiel :")
add_bullet("0-2 min : World's greatest stretch (fente avant + rotation du tronc) — 5 reps par côté")
add_bullet("2-4 min : Hip CARs — 5 reps par côté")
add_bullet("4-6 min : Knee to wall + calf raises lentes — 10 reps par côté")
add_bullet("6-8 min : Shoulder dislocates avec bâton — 10 reps")
add_bullet("8-10 min : T-spine rotation à 4 pattes — 10 reps par côté")
add_para("Tu peux faire cette routine à n'importe quel moment de la journée. Le meilleur moment : juste avant ton échauffement à l'entraînement, ou le matin pour bien commencer la journée.")

add_h2("Étirements après séance — quand et comment")
add_para("Après ta séance, tes muscles sont chauds, c'est le moment parfait pour faire des étirements légers ou utiliser un foam roller.")
add_para("Objectif : pas de gagner de l'amplitude (ça se fait à part), mais de relâcher les tissus tendus par la séance et favoriser la récupération.")
add_para("Comment faire :")
add_bullet("5 à 10 minutes après la séance")
add_bullet("Étirements LÉGERS sur les muscles principaux travaillés (jambes, dos, épaules selon ton type de séance)")
add_bullet("Pas de douleur. Si ça tire trop, recule.")
add_bullet("Tiens chaque position 30-45 secondes en respirant calmement")
add_bullet("Bonus : foam roller sur les zones tendues (mollets, ischios, quadriceps, fessiers, dos)")
add_para("Ne JAMAIS forcer un étirement sur un muscle blessé ou très fatigué. Le risque d'aggraver une micro-lésion est réel.")

add_h2("Pourquoi les étirements statiques AVANT l'entraînement = mauvaise idée")
add_para("C'est l'erreur que la majorité des sportifs font encore. Les étirements statiques (rester immobile dans une position d'étirement pendant 30s+) AVANT une séance, c'est l'inverse de ce que tu veux.")
add_para("La raison physiologique :")
add_bullet("Les étirements statiques relâchent les muscles et diminuent leur capacité à produire de la force rapidement")
add_bullet("Cet effet dure environ 1 heure")
add_bullet("Si tu fais des squats lourds, des sprints ou des sauts dans l'heure qui suit, tu performeras moins bien — et tu auras peut-être plus de risque de blessure (parce que le système de contrôle neuromusculaire est temporairement diminué)")
add_para("À la place, fais un échauffement dynamique :")
add_bullet("2-3 minutes de cardio léger (corde à sauter, vélo, footing léger)")
add_bullet("Mobilité articulaire (cercles épaules, hanches, chevilles)")
add_bullet("Mouvements dynamiques : skips, talons-fesses, montées de genoux, fentes marchées")
add_bullet("Accélérations progressives (3-5 sprints à 60%, 70%, 80% de ta vitesse max)")
add_para("Total : 10-15 minutes d'échauffement dynamique. Tes muscles sont chauds, ton système nerveux est activé, tu es prêt à performer.")
add_para("Règle simple : étirements statiques = APRÈS la séance. Échauffement dynamique = AVANT la séance. Jamais l'inverse.")
add_space()

add_h1(11, "Récupération")

add_h2("La récupération ne récompense pas celui qui s'entraîne le plus")
add_para("L'entraînement ne récompense pas celui qui en fait le plus. Il récompense celui qui récupère le mieux.")
add_para("À l'entraînement, tu crées un stimulus. Tu stresses le muscle. Tu imposes une adaptation.")
add_para("Mais la progression ne se fait pas pendant la séance. Elle se fait après.")
add_para("La croissance musculaire et le renforcement se produisent durant les heures et les jours où tu ne t'entraînes pas. Dépasser ton programme n'accélère pas tes résultats — au contraire, ça les freine.")
add_para("« Le surentraînement, c'est moins souvent du surentraînement qu'une mauvaise organisation. » — Fred Marcérou")
add_para("Autrement dit : les blocages de progression sont rarement liés à un manque de travail, mais presque toujours à une récupération mal calibrée.")

add_h2("La mécanique simple à comprendre")
add_bullet("Les muscles mettent en moyenne 48 à 72 heures à récupérer.")
add_bullet("Le système nerveux central met encore plus de temps.")
add_bullet("Sans récupération complète, ton intensité chute.")
add_bullet("Sans intensité, ta progression ralentit.")
add_para("C'est le principe de la surcompensation : tu travailles → tu fatigues → tu récupères → tu deviens plus fort. Si l'étape 3 est ratée, l'étape 4 n'arrive jamais.")

add_h2("Les 3 piliers de la récupération")

add_para("1. Le sommeil")
add_para("Minimum : 8 heures par nuit. Non négociable pour un athlète.")
add_para("C'est pendant le sommeil que :")
add_bullet("Tes fibres musculaires se reconstruisent")
add_bullet("Ton système nerveux se régénère")
add_bullet("Tes hormones (testostérone, hormone de croissance) se sécrètent")
add_para("Un bon sommeil = plus de force, plus d'intensité, plus de progression. Bâcler son sommeil pour s'entraîner plus, c'est l'erreur classique du sportif amateur.")

add_para("2. L'alimentation")
add_para("L'alimentation est ton carburant — comme l'essence pour une voiture. Sans apport suffisant :")
add_bullet("Pas de réparation musculaire optimale")
add_bullet("Pas de croissance")
add_bullet("Performances qui chutent à la séance suivante")
add_para("Détails complets dans le chapitre 8 (Nutrition athlète). Mais retiens une chose : après l'entraînement, tu as une fenêtre métabolique de 2 heures où ton corps assimile les nutriments 2 à 3 fois mieux. Profite-en.")

add_para("3. L'hydratation")
add_para("L'eau influence directement :")
add_bullet("Ta performance à l'entraînement et en match")
add_bullet("La contraction musculaire")
add_bullet("La récupération et la circulation sanguine")
add_para("Objectif minimum : 3 litres par jour pour un athlète actif. Plus si tu sues beaucoup (chaleur, séances longues, matches). Détails dans le chapitre 9 (Électrolytes).")
add_para("Règle simple : si tu as soif, tu es déjà déshydraté.")

add_h2("Le sommeil — comprendre ce qui le contrôle")

add_para("Les 2 moteurs de l'envie de dormir")
add_bullet("Le rythme circadien : ton horloge biologique. Elle crée des phases de fatigue et de pleine forme. Respecte-la — couche-toi à des heures régulières.")
add_bullet("La pression de sommeil (adénosine) : une molécule qui s'accumule dans ton cerveau tout au long de la journée. Au pic (souvent entre 14h et 17h), elle déclenche la sensation de sommeil.")

add_para("Le piège de la caféine")
add_para("La caféine bloque les récepteurs d'adénosine — tu te sens éveillé alors que l'adénosine continue de s'accumuler. Quand l'effet de la caféine disparaît, tu te crashes.")
add_para("Chiffres clés :")
add_bullet("Pic de concentration : 30 à 60 minutes après absorption")
add_bullet("Demi-vie : 5 à 7 heures (selon ta génétique)")
add_bullet("Exemple : un café à 17h → la moitié est encore active à minuit")
add_para("Règle simple : pas de caféine après 14h-15h si tu veux dormir correctement à 22h-23h.")
add_para("Dosage performance : 200 à 400 mg, 30 à 60 min avant l'entraînement. Au-delà : effet négatif sur la perf et la récupération.")

add_para("Le magnésium pour le sommeil")
add_para("Le magnésium a un effet calmant sur le système nerveux. Si tu en manques, ton sommeil devient agité (réveils fréquents, sensation de pas avoir dormi). 76% des Français sont carencés, beaucoup plus chez les athlètes qui suent.")
add_para("Protocole : 200-300 mg de magnésium (bisglycinate de préférence) 20 à 30 minutes avant le coucher.")

add_para("L'alcool — alerte rouge")
add_para("L'alcool dégrade ton sommeil de fond, ta récupération hormonale et ta composition corporelle. Un verre de vin ou une bière au dîner reste OK, mais l'abus tue tes performances.")
add_para("Si tu veux progresser sérieusement, limite à 2-3 verres par semaine maximum.")

add_h2("La récupération par l'alimentation — protocole simple")

add_para("La fenêtre métabolique (2h post-effort)")
add_para("Les 2 heures après ta séance sont les plus importantes pour la récup. Pourquoi ? Tes cellules musculaires sont ultra-sensibles à l'insuline, donc elles absorbent beaucoup plus facilement les acides aminés et le glucose. Résultat : reconstitution rapide du glycogène + démarrage de la synthèse protéique.")

add_para("Protocole 5 étapes autour de la séance")
add_bullet("EN CONTINU : assez de glucides dans ton alimentation quotidienne pour que tes stocks de glycogène soient pleins (le carburant de la performance).")
add_bullet("1 H AVANT : petite quantité de glucides simples (banane, dattes) + un peu de protéines pour activer le système nerveux.")
add_bullet("PENDANT (si > 1h) : 40 à 60 g de glucides par heure d'effort. Boisson maison ou fruits secs.")
add_bullet("JUSTE APRÈS (dans la 1ère heure) : 15-30 g de protéines + 45-90 g de glucides. Ratio 1/3 protéines pour 2/3 glucides.")
add_bullet("2 H APRÈS : repas solide complet (protéines + glucides + légumes) pour prolonger la récupération.")

add_para("Le shaker récupération — recette")
add_para("Recette utilisée par les joueurs pros du Toulouse Olympique XIII :")
add_bullet("1 banane mûre congelée (la maturité libère des enzymes qui restockent le glycogène + antioxydants)")
add_bullet("30 g de whey protéine")
add_bullet("1 cuillère à soupe de cacao (flavonoïdes, oxyde nitrique → meilleure oxygénation cellulaire)")
add_bullet("1 cuillère à café de curcuma (anti-inflammatoire puissant)")
add_bullet("1 pincée de poivre noir (la pipérine multiplie l'absorption du curcuma par ~20)")
add_bullet("1 cuillère à café de cannelle (anti-inflammatoire, probiotique)")
add_bullet("Beurre de cacahuète 100% naturel (bonnes graisses)")
add_bullet("Liquide : lait d'amande, lait de coco ou eau (pas de lait de vache en brique)")
add_para("À boire dans la première heure après ta séance.")

add_para("Glucides le soir = meilleur sommeil")
add_para("Consommer des glucides le soir augmente la sécrétion de sérotonine, qui se convertit en mélatonine (l'hormone du sommeil). Effets :")
add_bullet("Meilleur endormissement")
add_bullet("Sommeil plus profond")
add_bullet("Baisse du cortisol (l'hormone du stress, catabolique)")
add_para("Évite donc les régimes sans glucides le soir si tu es athlète. Le keto à long terme n'est pas adapté à la performance explosive.")

add_h2("Récupération active vs passive")

add_para("Tu as deux options pour gérer ta récup au jour le jour :")
add_para("Récupération PASSIVE = repos complet, pas d'activité physique. Idéale après une grosse séance ou un match. Le corps fait le boulot tout seul.")
add_para("Récupération ACTIVE = activité légère qui favorise la circulation sanguine sans stresser le système. Exemples :")
add_bullet("30-45 min de marche")
add_bullet("Vélo très tranquille")
add_bullet("Natation cool")
add_bullet("Mobilité + étirements doux (cf. chapitre 10)")
add_bullet("Foam roller")
add_para("La récupération active accélère l'élimination des déchets métaboliques et favorise le retour à la normale. Plus efficace qu'un canapé toute la journée — sauf si tu es vraiment cramé.")
add_para("Règle pratique : 1 séance de récup active de 30 min, le lendemain d'une grosse séance ou d'un match.")

add_h2("Les signaux d'overtraining à surveiller")

add_para("Si plusieurs de ces signes apparaissent en même temps et persistent, c'est que tu sur-sollicites ton corps :")
add_bullet("Fatigue qui s'accumule entre les séances et les semaines, sans descendre")
add_bullet("Récupération qui prend de plus en plus de temps malgré une charge identique")
add_bullet("Inflammations localisées qui persistent (tendons, articulations, muscles)")
add_bullet("Baisse progressive des performances (alors que ton entraînement n'a pas changé)")
add_bullet("Sommeil dégradé sans raison apparente")
add_bullet("Irritabilité, anxiété, motivation qui chute")
add_bullet("Système immunitaire affaibli (tu tombes plus souvent malade)")
add_bullet("Perte d'appétit")
add_bullet("Rythme cardiaque au repos plus élevé que d'habitude (5+ bpm)")
add_para("Si tu coches 3-4 de ces signes en même temps, tu es probablement en sur-sollicitation. À ajuster :")
add_bullet("Réduis le VOLUME (séries, distance, durée)")
add_bullet("Réduis l'INTENSITÉ (charge, vitesse, % effort max)")
add_bullet("Réduis la FRÉQUENCE (moins de séances/sem)")
add_bullet("Ajoute du sommeil et de la nutrition")

add_h2("Quand prendre un jour de repos supplémentaire")

add_para("La règle simple : si tu te poses la question, c'est probablement que tu en as besoin.")
add_para("Critères pour prendre un jour de repos non prévu :")
add_bullet("Tu as mal dormi 2 nuits de suite (< 6h)")
add_bullet("Tu as une charge mentale/stress externe importante (boulot, examens, vie perso)")
add_bullet("Tu sens des inflammations dans 2+ zones (genou + épaule, par exemple)")
add_bullet("Ta motivation à la séance est anormalement basse")
add_bullet("Ton rythme cardiaque au repos est 10+ bpm au-dessus de la normale")
add_para("Conseil : 1 jour de repos supplémentaire ne casse PAS ta progression. Au contraire, ça la sauve. Tenir à tout prix une séance que ton corps refuse = blessure ou stagnation.")
add_para("À l'inverse, si tu te sens bien, ne pas hésiter à pousser. L'équilibre se trouve dans l'écoute de soi, pas dans la rigidité d'un programme.")
add_space()

add_h1(12, "Mindset")

add_h2("Le vrai défi, c'est pas d'acheter des connaissances")
add_para("Si tu as acheté cet ebook, c'est que tu sais ce que tu veux : progresser, avancer, changer. Mais soyons honnêtes — si c'était facile, tout le monde aurait un corps de rêve et la discipline serait naturelle.")
add_para("Le vrai défi, ce n'est pas d'acheter des connaissances. C'est de les appliquer. De rester constant. De t'entraîner sérieusement, semaine après semaine, mois après mois. Ça demande du temps, de l'énergie, et surtout de la rigueur.")
add_para("Mais attention : trop d'effort sans équilibre, et tu craques. Pas assez, et les résultats ne viennent jamais. Alors comment trouver le juste milieu ?")

add_h2("Discipline > motivation")
add_para("La motivation est instable. Elle va et elle vient. Tu peux te sentir au top un lundi et complètement vidé un jeudi. C'est normal.")
add_para("La discipline, elle, doit rester.")
add_para("C'est précisément dans ces moments où tu n'as pas envie, où tu es fatigué, où tu pourrais facilement repousser à demain… que tout se joue. C'est là que tu fais la différence.")
add_para("La plupart des gens arrêtent à ce moment-là.")
add_para("Ne sois pas comme eux.")
add_para("Comme le dit Fred Marcérou : « La régularité est la clé, pas le No Pain No Gain. » Faire 3 séances correctes par semaine pendant 12 semaines vaut largement mieux que faire 5 séances héroïques avant de cramer.")

add_h2("Le continuum des motivations — pourquoi la tienne décide de ta longévité")
add_para("Il y a 6 niveaux de motivation, du plus puissant au plus fragile. Le sportif qui dure 10 ans n'est pas dans le même niveau que celui qui crame en 3 mois.")
add_para("Du plus fort au plus faible :")
add_bullet("MOTIVATION INTRINSÈQUE — Tu fais ton sport pour LUI-MÊME, par plaisir. Tu kiffes apprendre, tu kiffes te dépasser, tu kiffes les sensations. C'est la motivation la plus pérenne qui existe.")
add_bullet("RÉGULATION INTÉGRÉE — Le sport fait partie de qui tu ES. Tu cours parce que t'es quelqu'un qui court. Pas une activité, une identité. C'est ici qu'apparait \"l'identité d'athlète\" (voir plus bas).")
add_bullet("RÉGULATION IDENTIFIÉE — Tu comprends que c'est utile pour toi (santé, perf, mental). Tu acceptes l'effort parce que tu en vois la valeur.")
add_bullet("RÉGULATION INTROJECTÉE — Tu te forces parce qu'il \"faut absolument\". Pression interne, ego. Basé sur la peur de ne pas y arriver. Pas durable.")
add_bullet("RÉGULATION EXTERNE — Tu fais ça pour les likes Instagram, les compliments, ou par peur du jugement. Encore moins durable.")
add_bullet("AMOTIVATION — Tu ne vois plus le lien entre tes efforts et leurs effets. Zone rouge. Risque de burn-out, blessure, abandon.")
add_para("Le but du jeu : faire glisser ta motivation vers le HAUT du continuum. Quand le sport devient une partie de qui tu es (intégrée), la question \"j'y vais ou pas aujourd'hui ?\" disparaît. Tu y vas, point.")

add_h2("Comment gérer les baisses de motivation")
add_para("Quand la motivation chute (et ça arrivera), 4 trucs marchent :")
add_bullet("Réduis le seuil. Tu n'as pas envie de faire ta séance complète ? Promets-toi juste l'échauffement. 9 fois sur 10, tu finiras la séance. L'inertie de démarrer est plus dure que de continuer.")
add_bullet("Reconnecte-toi à ton pourquoi. Pourquoi t'es-tu mis à ça ? Quel objectif tu vises ? Si tu peux pas répondre, c'est que tes objectifs sont flous (relis le chapitre 2 : tu dois en avoir 3 minimum, écrits).")
add_bullet("Change le contexte. Séance impossible aujourd'hui ? Va à la salle quand même, fais 30 min de mobilité, mange ton repas perfect, hydrate-toi. La régularité gagne sur l'intensité.")
add_bullet("Lance-toi un challenge hebdomadaire. Battre tes reps à 80% du DC, finir un temps de rameur, faire X tractions. Un objectif court, atteignable, qui casse la routine. Marcérou recommande la fréquence HEBDO comme le bon compromis (le quotidien sature, le mensuel est trop lointain).")

add_h2("Sur-sollicitation vs sous-sollicitation — ne confonds pas")
add_para("Une baisse de motivation peut venir de 2 problèmes opposés. Mal identifier = aggraver.")
add_para("SUR-SOLLICITATION — tu en fais trop. Signaux :")
add_bullet("Fatigue qui ne descend pas entre séances")
add_bullet("Inflammations qui persistent")
add_bullet("Performances qui baissent malgré le même entraînement")
add_bullet("Motivation qui chute")
add_para("→ Solution : réduis VOLUME, INTENSITÉ ou FRÉQUENCE.")
add_para("SOUS-SOLLICITATION — tu en fais pas assez. Signaux :")
add_bullet("Tu ne progresses pas, mais tu ne fatigues jamais non plus")
add_bullet("Aucune inflammation, aucune sensation marquée")
add_bullet("Tu te trouves nul à l'entraînement, jamais au bord du seuil")
add_para("→ Solution : augmente VOLUME, INTENSITÉ ou FRÉQUENCE.")
add_para("Le piège : si t'es démotivé par manque de stimulation et que tu réduis encore les séances, tu accélères la chute. Si t'es démotivé par sur-sollicitation et que tu pousses plus fort, tu cramés. Diagnostique avant d'agir.")

add_h2("Recadrer la contrainte en opportunité")
add_para("La motivation chute souvent quand un imprévu arrive : blessure, vie pro chargée, examen, problème perso. La majorité abandonne à ce moment-là.")
add_para("L'athlète sérieux fait l'inverse : il REFORMULE l'objectif pour qu'il reste atteignable dans le contexte.")
add_para("Exemple : tu te blesses à un membre inférieur. Au lieu d'arrêter tout, tu pivotes vers :")
add_bullet("Travail respiratoire")
add_bullet("Renforcement haut du corps (bench, tirage, gainage)")
add_bullet("Mobilité, étirements profonds")
add_bullet("Lecture/formation sur ton sport")
add_para("Tu sors de la période de contrainte plus fort qu'avant, avec des angles que t'aurais jamais bossés en temps normal.")
add_para("Règle : la résilience, ce n'est pas \"continuer à faire la même chose envers et contre tout\". C'est \"reformuler intelligemment\".")

add_h2("Concentre-toi quand tu t'entraînes")
add_para("La salle, c'est un terrain de jeu. Mais un terrain de jeu SÉRIEUX.")
add_para("Pendant tes séries, sois concentré. Prépare-toi avant chaque répétition. Donne tout.")
add_para("Comme sur un match de tennis ou de foot : tu ne discutes pas pendant que tu joues.")
add_para("Si tu passes tes séances le téléphone à la main, à scroller entre les séries, à répondre à des messages, à discuter avec ton pote pendant 5 minutes entre chaque set — tu fais 50% de ta séance. Et tu auras 50% des résultats.")
add_para("Règle simple : 1h30 max à la salle. Téléphone en mode avion ou dans le sac. Concentration à 100%. Tu sors plus fort, pas plus fatigué.")

add_h2("Remplace les heures inutiles")
add_para("Regarde ce qui te vole du temps inutilement. Les heures à scroller sur les réseaux, à chercher la dose de dopamine facile.")
add_para("Transforme ce temps perdu en moments pour ton corps :")
add_bullet("Une séance de sport supplémentaire dans la semaine")
add_bullet("Du temps pour préparer tes repas (au lieu de bouffer de la merde par défaut)")
add_bullet("De la lecture sur ton sport ou ta nutrition")
add_bullet("Du sommeil — le levier que personne n'utilise vraiment")
add_para("Ces petits choix font la différence sur 6 mois. Pas sur 1 semaine. Sur 6 mois.")

add_h2("La technique habitudes — ajoute, ne supprime pas")
add_para("Pour faire émerger des comportements d'athlète, NE supprime PAS tes vieilles habitudes. AJOUTE des nouvelles jusqu'à ce qu'elles prennent le dessus.")
add_para("Pourquoi : le cerveau humain a beaucoup de mal avec la soustraction. En revanche, il est à l'aise avec l'addition, l'apprentissage.")
add_para("Exemples concrets :")
add_bullet("Au lieu de \"arrête de bouffer devant Netflix\" → ajoute \"mâche 15-20 fois chaque bouchée\" et \"pose la fourchette entre chaque\"")
add_bullet("Au lieu de \"vire le Coca\" → ajoute \"bois 500 ml d'eau avec une pincée de sel naturel au réveil\"")
add_bullet("Au lieu de \"arrête le scroll\" → ajoute \"lis 5 pages d'un livre par jour\"")
add_para("Les anciens comportements ne disparaissent pas par volonté. Ils disparaissent par substitution naturelle, parce que tu remplis ton temps autrement. C'est beaucoup plus efficace que la rigidité.")

add_h2("Identité d'athlète — ce que tu fais quand personne ne regarde")
add_para("C'est ce que tu fais quand personne ne regarde qui définit qui tu es vraiment.")
add_para("Tout le monde peut s'entraîner dur le jour d'un match. Tout le monde peut manger sain le jour où il poste sur Instagram. Tout le monde peut se coucher tôt la veille d'une compétition.")
add_para("L'athlète sérieux, c'est celui qui fait toutes ces choses en silence. Quand il n'y a rien à prouver. Quand il pourrait facilement faire autre chose. Quand personne ne saura jamais.")
add_para("Une identité, c'est pas une déclaration. C'est l'accumulation des choix invisibles que tu fais chaque jour.")
add_para("Pose-toi la question chaque matin : « Aujourd'hui, qu'est-ce que je fais qui prouve que je suis vraiment un athlète ? »")
add_para("Si la réponse est « rien », alors tu n'es pas un athlète ce jour-là. Tu es quelqu'un qui veut être un athlète. C'est pas la même chose.")

add_h2("Place du sport dans ta vie — l'introspection à faire")
add_para("Il n'y a pas de réponse universelle. Le sport peut être ta priorité numéro 1, ou il peut être derrière d'autres priorités (famille, boulot, études). Les deux sont valides.")
add_para("Mais tu dois te poser la question honnêtement :")
add_bullet("Quels sont les éléments de ta vie qui sont VRAIMENT importants pour toi ?")
add_bullet("Quelle place le sport doit-il avoir dans ce classement ?")
add_bullet("Es-tu prêt à dire NON à d'autres choses pour dire OUI à tes objectifs sportifs ?")
add_para("Une fois cette clarté installée, le reste devient simple. Tu sais quand tu peux pousser, quand tu dois te modérer, quand un choix est aligné avec ce que tu veux vraiment.")
add_para("Sans cette clarté, tu vas constamment hésiter, douter, abandonner. Avec elle, tu deviens implacable.")

add_h2("La règle qui résume tout")
add_para("Si tu vas à la salle pour t'entraîner — entraîne-toi vraiment.")
add_para("Si tu manges un repas pour ton corps — mange-le vraiment (pas en mode dégueulasse devant Netflix).")
add_para("Si tu dors pour récupérer — dors vraiment (pas le téléphone à 1h du matin).")
add_para("Tout ce que tu fais à moitié, tu le payes en résultats à moitié.")
add_space()

add_h1(13, "FAQ / Troubleshooting")
add_para("Toutes les questions qu'on me pose en DM. Si tu te poses une question qui n'est pas ici, la 17e question te dit comment me contacter.")

add_h2("Sur le programme")

add_para("1. À qui s'adresse ce programme ?")
add_para("À tous ceux qui veulent transformer leur physique, gagner en force, en masse musculaire, et améliorer leur performance sur le terrain. Que tu sois débutant ou sportif confirmé, il s'adapte (avec ajustements de charges et de volumes selon ton niveau).")

add_para("2. Combien de séances par semaine et combien de temps ?")
add_para("5 séances par semaine, environ 1h30 chacune. Échauffement + travail ciblé + récupération intégrée.")
add_para("Si tu peux pas faire 5 séances, vise 3-4. Mieux vaut 3 séances régulières pendant 12 semaines que 5 séances pendant 3 semaines avant de craquer.")

add_para("3. Que signifie « RIR » dans le programme ?")
add_para("RIR = Répétitions En Réserve. Ça indique combien de répétitions tu pourrais ENCORE faire après ta série.")
add_para("Exemple : 2 RIR = tu termines ta série en ayant encore 2 reps en réserve avant l'échec. C'est une façon plus précise de doser l'intensité que de juste dire « jusqu'à l'échec ».")

add_para("4. Quelle est la différence entre hypertrophie et force ?")
add_bullet("HYPERTROPHIE : développement de la masse musculaire. Charges modérées (70-85% du 1RM), beaucoup de répétitions, proximité de l'échec. C'est ce qui te fait grossir.")
add_bullet("FORCE : puissance pure, performance, stimulation du système nerveux. Charges lourdes (85-100% du 1RM), peu de répétitions. C'est ce qui te fait soulever lourd.")
add_para("Le programme combine les deux pour un développement complet d'athlète. Pas l'un OU l'autre, les deux.")

add_para("5. Le programme est-il sécurisé ?")
add_para("Oui, si tu respectes :")
add_bullet("Les charges adaptées à ton niveau (pas commencer trop lourd)")
add_bullet("Une bonne technique d'exécution (regarde les vidéos de chaque exo)")
add_bullet("Les temps de repos prescrits")
add_bullet("L'échauffement avant chaque séance")
add_para("Note importante : ce programme ne remplace pas un avis médical. Chaque utilisateur reste responsable de sa pratique. Si tu as des antécédents, consulte un médecin avant.")

add_para("6. Puis-je faire ce programme en pré-saison rugby ?")
add_para("Oui, c'est même son objectif principal. La pré-saison (~8 semaines) = LE bon moment pour suivre un programme structuré en blocs.")
add_para("Adaptation à faire selon ton calendrier : si tu as des entraînements terrain dans la semaine, réduis la fréquence de tes séances de muscu à 3-4/semaine pour gérer la fatigue globale. La règle : la performance sur le terrain prime sur la salle.")

add_h2("Sur la nutrition et la récup")

add_para("7. Ai-je besoin de compléments alimentaires pour progresser ?")
add_para("Non, ils sont optionnels. Ils servent à optimiser tes résultats et soutenir ta récupération. Mais une alimentation solide reste la priorité.")
add_para("Si tu veux investir, l'ordre des priorités (cf. chap 9) : créatine > oméga-3 > magnésium > vitamine D > sel naturel pour l'eau. Reste optionnel.")

add_para("8. Que contient le guide nutritionnel ?")
add_para("Le chapitre 8 te donne : calcul de tes besoins caloriques, choix d'objectif (prise de masse / perte gras / recompo), répartition des macronutriments, timing avant/pendant/après l'entraînement, timing autour des matches, exemple de journée 2800 kcal, et liste d'aliments sains + plaisir.")

add_para("9. Comment optimiser ma récupération et mon lifestyle ?")
add_para("Les 3 piliers (cf. chap 11) :")
add_bullet("Dors au moins 8h par nuit")
add_bullet("Mange assez de protéines et glucides (cf. chap 8)")
add_bullet("Hydrate-toi avec eau + sel naturel")
add_para("Bonus : écoute ton corps. La régularité et la constance sont tes meilleurs alliés. Mieux vaut 6 mois constant qu'un mois parfait suivi d'un crash.")

add_para("10. Puis-je suivre ce programme si je suis végétarien ou vegan ?")
add_para("Oui. Tu adaptes ton apport en protéines via :")
add_bullet("Tofu, tempeh, seitan")
add_bullet("Légumineuses (lentilles, pois chiches, haricots)")
add_bullet("Protéines végétales en poudre (pois, riz, chanvre)")
add_bullet("Quinoa, sarrasin")
add_bullet("Oléagineux (amandes, noix, graines)")
add_para("Conseil : tu auras besoin d'un peu plus de volume alimentaire pour atteindre tes 2,2-2,5 g de protéines par kg de masse maigre. Et complémentation B12 obligatoire en vegan.")

add_para("11. Es-ce que je peux faire du cardio en prise de masse ?")
add_para("Oui, c'est même recommandé, mais à faible intensité.")
add_para("L'objectif n'est pas de brûler un maximum de calories (sinon tu rentres en déficit et tu rates ta prise de masse), mais de :")
add_bullet("Maintenir ta santé cardio")
add_bullet("Améliorer ta récupération entre les séances de muscu")
add_bullet("Préserver tes capacités athlétiques")
add_para("Format type : 2-3 séances de 30-45 min de marche rapide, vélo tranquille ou natation. Pas de HIIT pendant une prise de masse sérieuse.")

add_h2("Troubleshooting — quand ça se passe pas comme prévu")

add_para("12. Que faire si je rate une séance ?")
add_para("Tu ne cherches PAS à la rattraper en doublant la suivante. Tu reprends le programme à la séance suivante normalement.")
add_para("Règle Marcérou : « La régularité est la clé, pas le rattrapage. »")
add_para("Pratique :")
add_bullet("1 séance manquée sur 3-4 : reprends normalement, ne modifie rien.")
add_bullet("2 séances manquées d'affilée : priorise la séance la plus importante du cycle (séance principale / jour lourd). Lâche les accessoires.")
add_bullet("3+ séances manquées : reviens à 1 semaine d'intro (intensité réduite, charges à 75-80% de tes max) avant de reprendre le bloc en cours.")

add_para("13. Si je peux pas faire un exercice (matériel manquant ou blessure mineure) ?")
add_para("Tu remplaces par un mouvement de la MÊME famille de pattern, en respectant l'intensité visée.")
add_para("Les patterns de mouvement :")
add_bullet("Genou : squat, leg press, fentes, step-ups, hack squat...")
add_bullet("Hanche : deadlift, romanian, hip thrust, good morning...")
add_bullet("Répulsion : DC, push press, dips, pompes...")
add_bullet("Tirage : tractions, rowing, face pull...")
add_bullet("Saut : CMJ, box jumps, broad jump...")
add_para("Exemples :")
add_bullet("Pas de rack pour back squat → goblet squat avec haltère, split squat, leg press")
add_bullet("Pas de barre pour DC → DC haltères, pompes lestées, dips")
add_para("Hiérarchie quand tu manques de temps ou de matos : main lifts (squat, DC, DT, OHP) prioritaires sur les accessoires. Si tu dois sacrifier, sacrifie depuis le bas.")
add_para("Pour une blessure : ne mets PAS tout en pause. Une blessure mineure à un membre laisse 80% de ton corps disponible. C'est l'occasion de bosser ce que tu négliges d'habitude (mobilité, gainage, travail respiratoire, technique sur autres patterns).")
add_para("⚠️ Une douleur qui persiste, qui s'aggrave ou qui change la qualité de mouvement → consulte un kiné. Pas d'adaptation maison.")

add_para("14. Si je suis fatigué : je force ou je repose ?")
add_para("Ni l'un ni l'autre par défaut. Tu identifies d'abord la nature de la fatigue.")
add_para("FATIGUE PONCTUELLE (mauvaise nuit, journée stressante) → tu adaptes la séance, tu ne l'annules pas :")
add_bullet("Réduis les top sets de 10-15%")
add_bullet("Passe de 4-5 séries à 3 séries")
add_bullet("Allonge les temps de récup entre séries")
add_bullet("Garde les main lifts, lâche les accessoires")
add_para("FATIGUE STRUCTURELLE (sur-sollicitation, signaux multiples sur plusieurs jours/semaines) → tu réduis ou tu reposes :")
add_bullet("Cf. chapitre 11 (Récupération) pour les signaux à reconnaître")
add_bullet("Quand prendre un jour de repos supplémentaire : si tu coches 3+ critères listés en fin de chapitre 11")
add_para("Le critical drop off point : si ta performance chute de plus de 7% en force ou 20% en hypertrophie en cours de séance, tu stop l'exercice. Tu ne creuses pas le trou.")

add_para("15. Combien de fois par semaine je peux faire des sprints ?")
add_bullet("HORS-SAISON : 2 séances de sprints par semaine, séparées de 48-72h.")
add_bullet("EN SAISON : 1 séance dédiée (le match du week-end compte déjà comme une grosse exposition sprint + changements de direction).")
add_para("Volume par séance :")
add_bullet("Accélération (0-30m) ou vitesse max (40-60m) : 100 à 300m total")
add_bullet("Capacité anaérobie alactique (80-150m) : doubler le volume")
add_bullet("Avec changements de direction (5-20m) : diviser par 2 (50-150m total)")
add_bullet("Avec charge additionnelle (gilet lesté, traineau) : diviser par 2")
add_para("Règle absolue : récup 1 minute par 10m de sprint. Sinon tu travailles la fatigue, pas la vitesse (cf. chap 4).")

add_para("16. À quel moment je teste à nouveau ma baseline ?")
add_para("Le tableau de tests baseline du chap 5 prévoit 3 mesures : Semaine 0 / Semaine 4 / Semaine 8.")
add_para("Pour le reste :")
add_bullet("Hebdomadaire : marqueurs perceptifs (sommeil, énergie, motivation, RPE des séances)")
add_bullet("Toutes les 4 semaines : recalibrer tes charges si tu sens que tu progresses (ou stagnes)")
add_bullet("Toutes les 8-12 semaines : test formel 1RM ou 3RM sur les main lifts")
add_para("Ne pas tester trop souvent : un test 1RM est coûteux nerveusement. C'est une séance entière qui n'apporte pas d'adaptation supplémentaire. Privilégie la méthode \"tester sans tester\" : tu observes l'évolution de tes RPE à charge constante, de tes reps à charge égale, de tes sauts (CMJ), etc.")

add_h2("Contact")

add_para("17. Puis-je te contacter pour des questions ou conseils ?")
add_para("Bien sûr. Tu me joins directement sur Instagram @rb_perform — je réponds aux questions pour t'aider à progresser. Sois précis dans ton message (contexte, ce que tu as essayé, ta question exacte) = je peux t'aider mieux et plus vite.")
add_space()

add_h1(14, "Conclusion + CTA")
for item in [
    "Récap des 8 semaines",
    "Et après ? Comment maintenir / passer au cycle suivant",
    "CTA principal : accompagnement perso (DM)",
    "CTA secondaire (si bundle actif) : Pack Saison Complète (Force et Masse + Athlète Explosif)"
]:
    add_bullet(item)
add_space()

add_h1(15, "Annexes")

add_h2("Annexe A — Tableau de tests baseline (à imprimer)")
add_para("Voir tableau du chapitre 5. Imprime cette page et garde-la à portée.")

add_h2("Annexe B — Planning hebdomadaire vide (8 semaines)")
add_para("À compléter pour visualiser ta progression sur les 8 semaines du programme.")

add_h2("Annexe C — Glossaire des termes techniques")
add_bullet("1RM : poids maximum que tu peux soulever sur 1 répétition.")
add_bullet("Hypertrophie : prise de masse musculaire.")
add_bullet("Concentrique : phase du mouvement où le muscle se contracte (= la remontée du squat).")
add_bullet("Excentrique : phase où le muscle s'étire sous charge (= la descente du squat).")
add_bullet("Pliométrie : entraînement basé sur des sauts/bondissements explosifs.")
add_bullet("RSA (Repeated Sprint Ability) : capacité à répéter des sprints courts avec peu de récup.")
add_bullet("VMA : Vitesse Maximale Aérobie.")
add_bullet("CMJ (Counter Movement Jump) : saut vertical avec contre-mouvement (descente puis saut).")

add_h2("Annexe D — Techniques d'intensification avancées (bonus)")
add_para("Pour ceux qui veulent intensifier leurs séances de musculation au-delà du programme de base.")

add_para("Séries dégressives")
add_para("Tu fais une série classique près de l'échec. Tu prends 15-20s de repos. Tu diminues la charge de 20-30% et tu enchaînes une nouvelle série jusqu'à l'échec. Garde la diminution la plus faible possible pour conserver de la tension mécanique tout en stimulant le muscle.")

add_para("Rest-pause")
add_para("Tu fais ta série principale. Tu reposes 15-20s SANS réduire la charge. Tu refais un maximum de répétitions avec le même poids. Tu récupères 30-40s. Tu refais une dernière série au maximum. Cette méthode pousse le muscle au-delà de la fatigue initiale.")

add_para("Tempo / négatifs lents")
add_para("Tu ralentis volontairement une des phases du mouvement — généralement l'excentrique (= la descente). Exemple : descente du squat en 4 secondes, remontée explosive. C'est sur la phase excentrique que les principaux dommages musculaires se produisent, donc c'est ce qui stimule le plus la croissance. Très efficace pour les débutants et intermédiaires.")

add_para("Supersets")
add_para("Tu enchaînes deux exercices pour des groupes musculaires antagonistes, sans repos entre les deux. Exemple classique : tractions (dos) + développé couché (pectoraux). Gain de temps + volume d'entraînement plus important en durée quasi identique.")

add_para("Bisets")
add_para("Comme les supersets mais sur deux groupes musculaires PAS forcément antagonistes. Exemple : tirage dos + curl biceps. Permet d'intensifier le travail d'une zone sans rallonger la séance.")

add_h2("Annexe E — Théorie athlète AVANCÉE (pour les nerds)")
add_para("Cette annexe va plus loin que le chapitre 4 pour ceux qui veulent les sources et les concepts précis utilisés en préparation physique de haut niveau.")

add_para("La nomenclature de Zatsiorski")
add_para("Trois types d'effort, classification proposée par Vladimir Zatsiorski (et utilisée par Fred Marcérou pour le rugby) :")
add_bullet("Efforts maximaux : 85-100% du 1RM, faible nombre de répétitions.")
add_bullet("Efforts dynamiques : ≤ 70% du 1RM avec intention de vitesse maximale.")
add_bullet("Efforts répétés : 70-85% du 1RM, beaucoup de répétitions jusqu'à proche de l'échec.")
add_para("La règle Louie Simmons + Fred Marcérou : « Quand l'athlète est faible, rends-le fort. Quand il est fort, rends-le rapide. Quand il est rapide, fais-le répéter. »")

add_para("Le CAT — Compensatory Acceleration Training (Dr. Hatfield)")
add_para("Principe : sur chaque répétition, tu pousses la charge avec une accélération maximale, peu importe le poids. González-Badilloa (2014) et Jones (1999) ont démontré des gains supérieurs en force et puissance avec le CAT vs un travail à vitesse réduite.")
add_para("Limite : difficile en groupe (concentration requise), et entraîne un freinage en fin de mouvement (= moins efficace sur le geste balistique). Solution : le VRT.")

add_para("Le VRT — Variable Resistance Training (élastiques)")
add_para("Tu ajoutes une résistance élastique progressive sur la barre. Effet : la charge augmente au fur et à mesure que tu remontes, ce qui force naturellement l'accélération maximale (le CAT devient automatique).")
add_para("Bénéfices documentés :")
add_bullet("Augmentation du RFD (Rate of Force Development).")
add_bullet("Maintien de l'accélération sur toute la phase concentrique.")
add_bullet("Stockage d'énergie élastique sur la phase excentrique (rappelle le triphasic de Cal Dietz).")
add_bullet("Réduction du risque articulaire (l'élastique freine, pas toi).")

add_para("La VBT — Velocity Based Training (capteurs de vitesse)")
add_para("Tu mesures la vitesse de la barre comme indicateur de performance. Exemple en développé couché à 80 kg :")
add_bullet("1RM = 100 kg → vitesse concentrique moyenne 0,45-0,55 m/s, 5-8 reps possibles.")
add_bullet("1RM = 120 kg → vitesse concentrique moyenne 0,60-0,75 m/s, 10-15 reps possibles.")
add_para("Conclusion : augmenter la Fmax a un impact direct sur la vitesse à charge sous-maximale et sur l'endurance de force. C'est l'argument fort pour le travail de force pure chez le rugbyman.")

add_para("Le triphasic training (Cal Dietz)")
add_para("Méthode qui découpe l'entraînement en 3 phases successives basées sur le cycle étirement-raccourcissement :")
add_bullet("Phase excentrique : ralentissement contrôlé pour maximiser le stockage d'énergie élastique.")
add_bullet("Phase isométrique : pause stricte au point bas pour habituer le système nerveux à la tension.")
add_bullet("Phase concentrique : explosivité maximale, intention de vitesse pure.")
add_para("L'objectif : raccourcir au maximum le temps de transition entre l'étirement et le raccourcissement = plus d'énergie élastique restituée = plus d'explosivité.")

add_para("Sources principales utilisées")
add_bullet("Zatsiorski V. — Science et pratique de l'entraînement de la force.")
add_bullet("Fred Marcérou — Préparation physique du rugby (Dragons Catalans).")
add_bullet("Cal Dietz — Triphasic Training (2012).")
add_bullet("Yuri Verkhoshansky — Pliométrie et entraînement explosif.")
add_bullet("Gilles Cometti — Méthodes modernes de musculation.")
add_bullet("Thierry Richard — Sport collectif et préparation physique.")
add_bullet("Charles Poliquin — Programmation et variations d'exercices.")
add_space()

# ─── Notes finales ──────────────────────────────────────────────────────────
add_h2("Notes structurelles importantes")
notes = [
    "Voix Rayan : direct, peer, expert, pas pédagogique lourd. Tutoiement partout (pas vouvoiement).",
    "Pas de mention 'coaching' ni 'individualisé' (raison APE + CQP ALS non validé). Préférer 'accompagnement', 'programme', 'prestation'.",
    "Cohérence DA avec la nouvelle charte RB Perform : fond dark, accent cyan, Inter Display (cover déjà refondue).",
    "CTAs distribués : 2-3 dans l'ebook (mi-livre, après baseline tests, fin), pas uniquement à la fin.",
]
for n in notes:
    add_bullet(n)

doc.save(OUTPUT)
print(f"Word : {OUTPUT}")
